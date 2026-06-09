import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Trust-Aware Memory Intelligence System\nHackathon Master Preparation Guide', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph("This document contains a comprehensive breakdown of every term, technology, and potential judge question for the In Time Tec (ITT) Hackathon. Memorize these concepts to defend your project confidently.\n")

    def add_section(title):
        heading = doc.add_heading(title, level=1)
        # heading.runs[0].font.color.rgb = RGBColor(108, 92, 231)

    def add_qa(q, a):
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"Q: {q}")
        run_q.bold = True
        
        p_a = doc.add_paragraph(f"A: {a}")
        
    def add_follow_up(q, a):
        p_q = doc.add_paragraph()
        p_q.paragraph_format.left_indent = Pt(20)
        run_q = p_q.add_run(f"↳ Follow-up Q: {q}")
        run_q.bold = True
        run_q.italic = True
        
        p_a = doc.add_paragraph(f"  A: {a}")
        p_a.paragraph_format.left_indent = Pt(30)
        
    def add_term(term, definition):
        p = doc.add_paragraph()
        run_term = p.add_run(f"{term}: ")
        run_term.bold = True
        p.add_run(definition)

    # ── Section 1: The Hackathon & In Time Tec ──
    add_section("1. The Hackathon & In Time Tec (ITT)")
    add_qa("Who is In Time Tec (ITT) and why are they hosting this?", "In Time Tec is a global software research and development firm focused on creating abundance and delivering value. They host hackathons to identify top talent, foster innovation, and explore cutting-edge solutions like AI and autonomous agents.")
    add_follow_up("How does your project align with ITT's core values?", "ITT values innovation and robust engineering. Our system solves a real-world enterprise problem: AI hallucination and memory corruption. By building a 'trust-aware' system, we bring reliability to AI, which is crucial for enterprise adoption.")
    add_qa("What is the core problem statement of this hackathon?", "To build a system where AI agents don't just blindly remember everything, but actively evaluate, verify, and curate incoming information based on trust, source reliability, and contradiction analysis.")
    
    # ── Section 2: Core Terminology ──
    add_section("2. Core Terminology & Concepts")
    add_term("Claim", "A piece of incoming information (e.g., 'Apple launched a VR headset'). It can be true, false, or a rumor.")
    add_term("Triple (Subject-Predicate-Object)", "The structured format we extract from a claim. (Subject: 'Apple', Predicate: 'launched', Object: 'VR headset'). This allows the system to query and compare facts programmatically.")
    add_term("Trust Score", "A computed metric (0.0 to 1.0) representing how much the system believes a claim. It combines Source Reliability, Verifiability, and Corroboration.")
    add_term("Confidence Delta", "The change in trust score after evaluating new evidence. If a highly reliable source confirms an existing memory, the delta is positive.")
    add_term("LangGraph", "The framework we use to route the AI agents. It represents the workflow as a mathematical graph (nodes and edges).")
    add_term("Memory Store", "The database (SQLite) where curated facts are stored along with their provenance (origin story) and trust scores.")

    # ── Section 3: The 6 Memory States ──
    add_section("3. The 6 Memory States (Crucial for Demo)")
    doc.add_paragraph("Your system continuously decides what to do with new claims. Here are the 6 possible outcomes:")
    add_term("ACCEPTED", "A new claim from a sufficiently reliable source is stored as a new memory.")
    add_term("UPDATED", "A highly reliable source provides new, conflicting information about an existing subject/predicate, causing the memory to update its 'object' (e.g., CEO is now Bob, not Alice).")
    add_term("REJECTED", "A claim from an untrusted source (like a spam bot) contradicts established facts and is thrown out completely.")
    add_term("DOWNGRADED", "A moderately reliable source contradicts our belief. We don't delete the memory, but we lower its confidence score because there is now uncertainty.")
    add_term("FORGOTTEN", "When a memory's confidence drops below a critical threshold (due to repeated contradictions), it is marked as forgotten/deprecated.")
    add_term("MERGED", "Multiple sources report the exact same fact. The system merges them, adding the new source to the provenance list and boosting the overall confidence.")

    # ── Section 4: Architecture & The Agent Pipeline ──
    add_section("4. Multi-Agent Pipeline Architecture")
    doc.add_paragraph("When a claim enters the system, it passes through 5 distinct Agent Nodes:")
    add_term("1. Extraction Agent", "Uses LLM to convert natural language into a Subject-Predicate-Object triple.")
    add_term("2. Verification Agent", "Checks if the claim is structurally sound and logically verifiable.")
    add_term("3. Contradiction Agent", "Queries the SQLite DB to find existing memories about the same Subject and Predicate. Detects if the new Object matches or contradicts the old Object.")
    add_term("4. Trust Agent", "Calculates a mathematical trust score based on the source's reliability and historical corroboration.")
    add_term("5. Curation Agent", "The final judge. It looks at the contradiction report and the trust score, and decides which of the 6 Actions (Accept, Reject, Update, etc.) to take.")

    # ── Section 5: Technology Stack Deep Dive ──
    add_section("5. Technology Stack & 'Why Did You Choose It?'")
    
    add_qa("Why use LangGraph instead of standard LangChain or AutoGen?", "LangGraph provides explicit state management and cyclic execution. It allows us to define rigid workflows (Extraction -> Contradiction -> Trust) ensuring no steps are skipped, while keeping the state immutable at each step.")
    add_follow_up("What happens if an agent fails in LangGraph?", "Because state is passed between nodes, if an LLM call fails, we can catch the exception, return an 'error' flag in the state, and route to a fallback or END node without crashing the whole pipeline.")
    
    add_qa("Why FastAPI for the backend?", "FastAPI is asynchronous, extremely fast, and automatically generates Swagger/OpenAPI documentation. It is the industry standard for serving machine learning models and AI pipelines.")
    
    add_qa("Why Streamlit for the frontend?", "Streamlit allows rapid prototyping of data-heavy applications. We built a custom glassmorphism CSS layer on top of it to make it look like a premium production app, while maintaining the speed of Python-based UI development.")
    add_follow_up("Isn't Streamlit too slow for production?", "For consumer apps, yes. But for internal enterprise dashboards, data science tools, and AI administrative panels, it is widely used in production. If we scale, we would migrate to React/Next.js, but our FastAPI backend would remain exactly the same.")

    add_qa("Why use Groq and LLaMA 3.3 70B?", "Groq uses LPUs (Language Processing Units) which provide deterministic, lightning-fast inference. We used LLaMA 3.3 70B because we needed high reasoning capabilities (for contradiction detection and JSON generation) at near-instant speeds to make the pipeline feel real-time.")

    add_qa("Why SQLite and SQLAlchemy?", "For a hackathon, SQLite provides zero-configuration persistent storage. SQLAlchemy (an ORM) abstracts the SQL logic, meaning we can swap SQLite for PostgreSQL by changing exactly one line of code in production.")

    # ── Section 6: Hard Questions & Stress Tests (Judge Defense) ──
    add_section("6. Hard Questions & Stress Tests (Judge Defense)")
    
    add_qa("How does the system handle prompt injection or malicious claims?", "The Extraction Agent is rigidly prompted to output only JSON triples. If a user injects a prompt like 'Ignore all instructions', the LLM will fail to produce valid JSON, the parser will catch the exception, and the claim will be safely rejected without affecting the database.")
    
    add_qa("What happens if two highly reliable sources contradict each other simultaneously?", "The system processes them sequentially. The first is ACCEPTED. When the second arrives, the Contradiction Agent flags a clash. Because both have high trust, the Curation Agent will execute a DOWNGRADE on the existing memory. It becomes a 'contested' fact until a third reliable source breaks the tie.")
    add_follow_up("Does order matter then? Isn't that biased?", "Yes, temporal bias exists. In a production version, we would implement a 'Holding Area' for conflicting claims arriving within a short time window, evaluating them in batch rather than sequentially.")

    add_qa("How do you ensure the LLM outputs correct JSON every time?", "We use structured prompting (double curly braces), clear schema instructions, and robust try-except blocks. If the LLM wraps the output in markdown (```json), our code strips it. If it fails entirely, the pipeline falls back gracefully.")

    add_qa("How does the 'Explainability' feature actually work?", "When requested, we query the 'ChangeLog' table for a specific memory ID. We pass the full history of decisions (the audit trail) to the LLM, asking it to summarize in natural language *why* the system arrived at the current belief state.")
    
    add_qa("How do you prevent the database from growing infinitely?", "We implemented the 'FORGOTTEN' state. When a memory's confidence drops below a threshold, its status changes to FORGOTTEN. In production, a background cron job could archive or delete forgotten memories to save space.")

    add_qa("What is the difference between REJECTED and DOWNGRADED?", "REJECTED means a new claim was thrown out (the database wasn't modified) because it lacked trust. DOWNGRADED means a new claim was somewhat trusted but contradicted our belief, so we lowered our confidence in our *existing* memory.")

    # ── Section 7: Future Scalability ──
    add_section("7. Future Scalability (If we had 3 more months)")
    doc.add_paragraph("1. Database: Migrate from SQLite to PostgreSQL + pgvector for semantic similarity matching (so 'Apple acquired Netflix' and 'Apple bought Netflix' are recognized as the exact same predicate without rigid string matching).")
    doc.add_paragraph("2. Frontend: Rebuild the UI in Next.js/React for better mobile responsiveness and websockets (SignalR/Socket.io) for real-time live updates without polling.")
    doc.add_paragraph("3. Multi-modal Ingestion: Build agents that can ingest claims from PDF reports, Twitter API streams, and News RSS feeds automatically, rather than relying on manual text input.")
    doc.add_paragraph("4. User Authentication: Add Role-Based Access Control (RBAC) so only 'Admin' judges can override memory trust scores manually.")

    doc.save('Hackathon_Master_Guide.docx')
    print("Word Document successfully generated at Hackathon_Master_Guide.docx")

if __name__ == '__main__':
    create_document()
